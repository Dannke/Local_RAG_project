"""Load source documents from local files."""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from rag_project.models import Document


def load_text_files(
    input_dir: str | Path,
    patterns: Iterable[str] = ("*.txt", "*.md"),
    encoding: str = "utf-8",
) -> list[Document]:
    root = Path(input_dir)
    documents: list[Document] = []

    for pattern in patterns:
        for path in sorted(root.rglob(pattern)):
            if not path.is_file():
                continue
            text = path.read_text(encoding=encoding)
            relative_path = path.relative_to(root).as_posix()
            documents.append(
                Document(
                    id=relative_path,
                    text=text,
                    metadata={
                        "source": str(path),
                        "relative_path": relative_path,
                    },
                )
            )

    return documents


def load_documents(
    input_dir: str | Path,
    patterns: Iterable[str] = ("*.txt", "*.md", "*.pdf", "*.docx"),
    encoding: str = "utf-8",
) -> list[Document]:
    root = Path(input_dir)
    documents: list[Document] = []

    for pattern in patterns:
        for path in sorted(root.rglob(pattern)):
            if not path.is_file():
                continue
            documents.extend(load_document_file(path, root=root, encoding=encoding))

    return documents


def load_document_file(
    path: str | Path,
    root: str | Path | None = None,
    encoding: str = "utf-8",
) -> list[Document]:
    file_path = Path(path)
    root_path = Path(root) if root is not None else file_path.parent
    relative_path = file_path.relative_to(root_path).as_posix()
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf_documents(file_path, relative_path)

    text = extract_text(file_path, encoding=encoding)
    if not text.strip():
        return []

    metadata = {
        "source": str(file_path),
        "relative_path": relative_path,
        "file_type": suffix.lstrip("."),
    }
    if suffix == ".docx":
        metadata["page"] = 1
        metadata["page_label"] = "1"

    return [Document(id=relative_path, text=text, metadata=metadata)]


def extract_text(path: str | Path, encoding: str = "utf-8") -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding=encoding)
    if suffix == ".pdf":
        return _extract_pdf_text(file_path)
    if suffix == ".docx":
        return extract_docx_text(file_path)

    raise ValueError(f"Unsupported document type: {file_path.suffix}")


def _extract_pdf_text(path: Path) -> str:
    return "\n\n".join(document.text for document in load_pdf_documents(path, path.name))


def load_pdf_documents(path: str | Path, relative_path: str | None = None) -> list[Document]:
    file_path = Path(path)
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Install pypdf to load PDF files: python -m pip install pypdf") from exc

    reader = PdfReader(str(file_path))
    source_name = relative_path or file_path.name
    documents: list[Document] = []

    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        documents.append(
            Document(
                id=f"{source_name}:page-{index:04d}",
                text=text,
                metadata={
                    "source": str(file_path),
                    "relative_path": source_name,
                    "file_type": "pdf",
                    "page": index,
                    "page_label": str(index),
                },
            )
        )

    return documents


def extract_docx_text(path: str | Path) -> str:
    """Extract paragraph text from a DOCX file."""

    file_path = Path(path)
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        # Provide more helpful error message with debugging info
        import sys
        error_msg = (
            f"Failed to import python-docx. "
            f"Install it with: python -m pip install python-docx\n"
            f"Python: {sys.executable}\n"
            f"Error: {exc}"
        )
        raise RuntimeError(error_msg) from exc

    try:
        doc = DocxDocument(str(file_path))
        paragraphs = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ]
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Error processing DOCX file {file_path}: {e}") from e
