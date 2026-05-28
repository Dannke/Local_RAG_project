import tempfile
import unittest
from pathlib import Path

from docx import Document as DocxDocument

from rag_project.ingestion.loaders import extract_docx_text, load_documents


class DocxLoaderTests(unittest.TestCase):
    def test_extract_docx_text_reads_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "example.docx"
            doc = DocxDocument()
            doc.add_paragraph("First paragraph.")
            doc.add_paragraph("")
            doc.add_paragraph("Second paragraph.")
            doc.save(path)

            text = extract_docx_text(path)

            self.assertEqual(text, "First paragraph.\n\nSecond paragraph.")

    def test_load_documents_includes_docx_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "docs" / "example.docx"
            path.parent.mkdir()
            doc = DocxDocument()
            doc.add_paragraph("DOCX content.")
            doc.save(path)

            documents = load_documents(temp_dir, patterns=("*.docx",))

            self.assertEqual(len(documents), 1)
            self.assertEqual(documents[0].text, "DOCX content.")
            self.assertEqual(documents[0].metadata["file_type"], "docx")
            self.assertEqual(documents[0].metadata["relative_path"], "docs/example.docx")
            self.assertEqual(documents[0].metadata["page"], 1)
